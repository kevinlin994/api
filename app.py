from flask import Flask, request, redirect, url_for, flash, jsonify, send_file, render_template
from werkzeug.exceptions import Forbidden, HTTPException, NotFound, RequestTimeout, Unauthorized
import numpy as np
import pandas as pd
import pickle as p
import json
from summarization.textsummarization import bert_sum
from docx import Document
from docx.shared import Inches
from linkcheck import flag_private_urls

app = Flask(__name__)

### Testing out HTML
@app.route('/')
def form():
  return render_template('index.html')

@app.errorhandler(NotFound)
def page_not_found_handler(e: HTTPException):
    return render_template('404.html'), 404


@app.errorhandler(Unauthorized)
def unauthorized_handler(e: HTTPException):
    return render_template('401.html'), 401


@app.errorhandler(Forbidden)
def forbidden_handler(e: HTTPException):
    return render_template('403.html'), 403


@app.errorhandler(RequestTimeout)
def request_timeout_handler(e: HTTPException):
    return render_template('408.html'), 408
### THIS ONE WORKS WITH FORM-DATA
@app.route('/api/v1/resources/text/', methods=['POST'])
def summarize_from_text():
    
    data = request.form["data"]
    summary = bert_sum(data)

    return summary

### THIS ONE WORKS WITH DOCX FILES
@app.route('/api/v1/resources/document/summary', methods=['GET', 'POST'])
def summarize_from_file():
    
    f = request.files['data']
    f.save('datafile.docx')
    document = Document('datafile.docx')
    text =''
    for para in document.paragraphs:
        text+=para.text
    
    summary = bert_sum(text)

    return summary

### THIS ONE WORKS WITH DOCX FILES
@app.route('/api/v1/resources/document/docx', methods=['GET', 'POST'])
def return_document():
    f = request.files['data']
    f.save('datafile.docx')

    return send_file('datafile.docx', attachment_filename='test.docx')

def transform():
    data = 'datafile.docx'
    document = Document(data)
    new = Document()
    text =''
    for para in document.paragraphs:
        text+=para.text
    new.add_paragraph(bert_sum(text))
    return new

### Allows for HTML interface
@app.route('/transform', methods=["GET","POST"])
def transform_view():
    f = request.files['data_file']
    f.save('datafile.docx')
    if not f:
        return "Please upload a word document"
    result = transform()
    result.save('result.docx')
    return send_file('result.docx', attachment_filename='new_file.docx')

### THIS ONE WORKS WITH FORM-DATA
@app.route('/api/v1/resources/document/links', methods=['POST'])
def check_links():
    
    data = request.files["link"]
    data.save('linkfile.docx')
    results = flag_private_urls('linkfile.docx')
    results.to_excel('links.xlsx')
    
    return send_file('links.xlsx')

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')
